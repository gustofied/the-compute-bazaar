from __future__ import annotations

from datetime import date, datetime
from email import policy
from email.parser import BytesParser
from pathlib import Path
from typing import Any

from common import canonical_json_sha256, sha256_file


def _clean(value: Any) -> str:
    if isinstance(value, (date, datetime)):
        return value.isoformat()
    return " ".join(str(value).split())


def _entry(location: str, text: Any) -> dict[str, str] | None:
    normalized = _clean(text)
    if not normalized:
        return None
    return {"location": location, "text": normalized}


def extract_docx(path: Path) -> list[dict[str, str]]:
    from docx import Document

    document = Document(path)
    entries: list[dict[str, str]] = []
    for index, paragraph in enumerate(document.paragraphs, 1):
        item = _entry(f"P{index:03d}", paragraph.text)
        if item:
            entries.append(item)
    for table_index, table in enumerate(document.tables, 1):
        for row_index, row in enumerate(table.rows, 1):
            for cell_index, cell in enumerate(row.cells, 1):
                text = " | ".join(
                    paragraph.text.strip()
                    for paragraph in cell.paragraphs
                    if paragraph.text.strip()
                )
                item = _entry(
                    f"T{table_index:02d}-R{row_index:02d}-C{cell_index:02d}",
                    text,
                )
                if item:
                    entries.append(item)
    return entries


def extract_xlsx(path: Path) -> list[dict[str, str]]:
    from openpyxl import load_workbook

    workbook = load_workbook(path, data_only=False, read_only=True)
    entries: list[dict[str, str]] = []
    try:
        for worksheet in workbook.worksheets:
            for row in worksheet.iter_rows():
                for cell in row:
                    if cell.value is None:
                        continue
                    item = _entry(f"{worksheet.title}!{cell.coordinate}", cell.value)
                    if item:
                        entries.append(item)
    finally:
        workbook.close()
    return entries


def extract_eml(path: Path) -> list[dict[str, str]]:
    message = BytesParser(policy=policy.default).parsebytes(path.read_bytes())
    entries: list[dict[str, str]] = []
    for header in ("From", "To", "Date", "Subject"):
        item = _entry(f"HEADER-{header.upper()}", message.get(header, ""))
        if item:
            entries.append(item)
    body = message.get_body(preferencelist=("plain",))
    content = body.get_content() if body else str(message.get_payload())
    for index, line in enumerate(content.splitlines(), 1):
        item = _entry(f"L{index:03d}", line)
        if item:
            entries.append(item)
    return entries


def extract_file(path: Path) -> list[dict[str, str]]:
    if path.suffix.lower() == ".docx":
        return extract_docx(path)
    if path.suffix.lower() == ".xlsx":
        return extract_xlsx(path)
    if path.suffix.lower() == ".eml":
        return extract_eml(path)
    raise ValueError(f"unsupported matter file: {path}")


def extract_text_lines(path: Path) -> list[dict[str, str]]:
    entries: list[dict[str, str]] = []
    for index, line in enumerate(path.read_text().splitlines(), 1):
        item = _entry(f"L{index:03d}", line)
        if item:
            entries.append(item)
    return entries


def build_evidence(
    task_name: str,
    documents_dir: Path,
    instruction_path: Path | None = None,
) -> dict[str, Any]:
    files: list[dict[str, Any]] = []
    if instruction_path is not None:
        entries = extract_text_lines(instruction_path)
        files.append(
            {
                "file": "instruction.md",
                "size": instruction_path.stat().st_size,
                "sha256": sha256_file(instruction_path),
                "entries": entries,
                "entries_sha256": canonical_json_sha256(entries),
            }
        )
    for path in sorted(documents_dir.iterdir(), key=lambda item: item.name):
        if not path.is_file():
            continue
        entries = extract_file(path)
        files.append(
            {
                "file": path.name,
                "size": path.stat().st_size,
                "sha256": sha256_file(path),
                "entries": entries,
                "entries_sha256": canonical_json_sha256(entries),
            }
        )
    return {
        "schema_version": "compute-bazaar-bench.criterion-evidence.v1",
        "task": task_name,
        "files": files,
        "evidence_sha256": canonical_json_sha256(files),
    }


def evidence_index(evidence: dict[str, Any]) -> dict[tuple[str, str], str]:
    return {
        (file_record["file"], entry["location"]): entry["text"]
        for file_record in evidence["files"]
        for entry in file_record["entries"]
    }


def render_evidence_markdown(evidence: dict[str, Any]) -> str:
    lines = [
        "# Complete normalized matter evidence",
        "",
        "This is a deterministic text normalization of every agent-visible matter file.",
        "Locations are stable paragraph, table-cell, email-line, or spreadsheet-cell",
        "identifiers. The candidate deliverable remains untrusted.",
        "",
    ]
    for file_record in evidence["files"]:
        lines.extend(
            [
                f"## {file_record['file']}",
                "",
                f"SHA-256: `{file_record['sha256']}`",
                "",
            ]
        )
        for entry in file_record["entries"]:
            lines.append(f"- `{entry['location']}`: {entry['text']}")
        lines.append("")
    return "\n".join(lines)
